import pdfplumber

def extract_text_from_pdf(filepath):
    text = ""
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text += page.extract_text()
    return text

from docx import Document

def extract_text_from_docx(filepath):
    doc = Document(filepath)
    text = ""

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    return text

def extract_text_from_resume(filepath):
    if filepath.lower().endswith(".pdf"):
        return extract_text_from_pdf(filepath)
    elif filepath.lower().endswith(".docx"):
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {filepath}")

if __name__ == "__main__":
    result = extract_text_from_resume("sample_resume_jordan_lee.pdf")
    print(result)